# -*- coding: utf-8 -*-
import streamlit as st
from dataclasses import dataclass
from typing import Dict, Any, List
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.lib.units import mm
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI  # DeepSeek uses OpenAI-compatible SDK

st.set_page_config(page_title="石家庄装备制造产业园 投决会研判（严格口径 + DeepSeek）", layout="wide")

# ---------------- 配置 ----------------
CONFIG = {
    "thresholds": {"investPerMu": 300.0, "taxPerMu": 25.0},
}
INDUSTRIES = {"low": "低空经济", "svc": "服务类", "eqp": "装备制造类"}
PROJECT_TYPES = {
    "land": "征地类型",
    "existingNoPolicy": "购买或者租赁园区内的商业或园区自有产业厂房（不需要政策支持）",
    "ownFactoryWithPolicy": "购买园区自有的厂房（需要政策支持）",
}
NEED_TYPES = {
    "buy": "购买厂房",
    "rent": "租赁厂房",
    "ipark": "产业港定制建设",
    "buy_land": "购买土地",
}
CARRIERS = {"kcg": "科创谷厂房", "ipark": "产业港厂房", "social": "社会现房/园区内商业"}

# ---------------- 工具 ----------------
def compute_mu(land_mu: float, building_area: float, floor_ratio: float) -> float:
    mu_by_land = land_mu if land_mu else 0.0
    mu_by_build = (building_area / (floor_ratio * 666.67)) if (building_area and floor_ratio) else 0.0
    return mu_by_land if mu_by_land > 0 else (mu_by_build if mu_by_build > 0 else 0.0)

def fnum(v, digits=0):
    try: return f"{float(v):.{digits}f}"
    except: return str(v)

# ---------------- 数据 ----------------
@dataclass
class Model:
    # 简介
    projectName: str = ""
    investTotal: float = 0.0
    locate: str = ""
    landMu: float = 0.0
    buildingArea: float = 0.0
    floorRatio: float = 0.0
    introContent: str = ""
    expectedOutput: float = 0.0
    expectedAnnualTax: float = 0.0
    expectedJobs: int = 0

    # 研判
    industry: str = "eqp"
    companyName: str = ""
    establishedYear: str = ""
    registeredAt: str = ""
    isLuanReg: bool = True
    importBusiness: str = ""
    newBusiness: str = ""

    needType: str = "buy"
    carrier: str = "kcg"
    techTitles: str = ""
    chainMaturity: str = "成熟"
    innovation: str = "较强"
    customerStable: str = "稳定"
    marketBase: str = "扎实"
    chainSegmentFill: str = ""

    revenueY2: float = 0.0
    revenueY1: float = 0.0
    taxY2: float = 0.0
    taxY1: float = 0.0
    industryTrend: str = "向好"

    # 风险与意向
    riskDishonest: bool = False
    riskEnv: bool = False
    riskIllegalLand: bool = False
    riskLicenseMissing: bool = False
    intentAgree: bool = True

def evaluate(m: Model) -> Dict[str, Any]:
    mu = compute_mu(m.landMu, m.buildingArea, m.floorRatio)
    invest_intensity = (m.investTotal / mu) if mu > 0 else 0.0
    tax_intensity = (m.expectedAnnualTax / mu) if mu > 0 else 0.0
    thI = CONFIG["thresholds"]["investPerMu"]
    thT = CONFIG["thresholds"]["taxPerMu"]
    pass_hard = (invest_intensity >= thI and tax_intensity >= thT)
    invest_need = max(0.0, thI * mu - m.investTotal)
    tax_need = max(0.0, thT * mu - m.expectedAnnualTax)

    veto_reasons = []
    if m.riskDishonest: veto_reasons.append("失信被执行/严重信用风险")
    if m.riskEnv: veto_reasons.append("重大环保/安监处罚未结")
    if m.riskIllegalLand: veto_reasons.append("违法违规用地")
    if m.riskLicenseMissing: veto_reasons.append("核心资质缺失且短期不可补齐")
    veto = len(veto_reasons) > 0

    if veto:
        decision = "暂缓/拒绝"
    elif pass_hard:
        decision = "通过/签约"
    else:
        decision = "附条件通过"

    return {
        "mu": mu,
        "investIntensity": invest_intensity,
        "taxIntensity": tax_intensity,
        "passHard": pass_hard,
        "investNeed": invest_need,
        "taxNeed": tax_need,
        "decision": decision,
        "veto": veto,
        "vetoReasons": veto_reasons,
        "thI": thI, "thT": thT,
    }

# ---------------- 报告文本（严格按原文） ----------------
def build_text(m: Model, ev: Dict[str, Any]) -> Dict[str, str]:
    intro = (
        f"{m.projectName}+计划投资{fnum(m.investTotal,0)}万元，拟选址{m.locate}，"
        f"占地{fnum(m.landMu,2)}亩/实际{fnum(m.buildingArea,0)}平米，"
        f"建设内容：{m.introContent}，预计经济效益（预计年产值{fnum(m.expectedOutput,0)}万元、"
        f"预计年税收{fnum(m.expectedAnnualTax,0)}万元、预计带动就业人数{m.expectedJobs}人）。"
    )

    if m.isLuanReg:
        reg_text = f"拟将{m.importBusiness}业务导入园区" + (f"，拓展{m.newBusiness}新业务" if m.newBusiness else "") + "。"
    else:
        reg_text = f"尚未注册于园区，拟将{m.importBusiness}业务导入园区" + (f"，拓展{m.newBusiness}新业务" if m.newBusiness else "") + "。"
    judge1 = (
        f"项目是什么项目（园区的产业大类主要是低空经济、服务类、装备制造类），"
        f"符合园区{INDUSTRIES[m.industry]}产业的发展规划，项目主体是{m.companyName}，"
        f"{m.establishedYear}年注册于{m.registeredAt}，{reg_text}"
    )

    need_phrase = (
        "该项目的主要需求是购买厂房还是租赁/购买厂房（厂房有园区自有的科创谷的厂房，"
        "也有正在建设的产业港的厂房产业港的厂房就是企业提出需求，园区在园区自有的地块上根据企业的需求进行建设。）"
        "然后就是项目开展什么业务，"
    )
    ability_phrase = (
        f"项目主体是不是已经稳定运行多年了，有无技术或者河北省乃至国家的相关称号（{m.techTitles or '—'}），"
        f"是否具有成熟的产业链条和完善的技术体系，技术创新能力{m.innovation}，"
        f"是否具有稳定的客户资源和扎实的市场基础（{m.customerStable}/{m.marketBase}），"
        f"企业的业务部分在入园之后能够填补相关产业链的{m.chainSegmentFill or '—'}环节，增强园区在什么环节的实力。"
        f"（本项目拟开展业务：{m.introContent or '—'}；主要需求：{NEED_TYPES.get(m.needType,'')}；拟承接载体：{CARRIERS.get(m.carrier,'')}。）"
    )
    judge2 = need_phrase + ability_phrase

    judge3 = (
        "企业的近两年的营收和税收多少，是不是稳中向好的，这个业务目前的发展趋势是什么，"
        "在这个发展趋势下，企业的经济效益是不是向好的，如果目前企业已经有经济效益了就可以看一下近两年的经济效益是不是在上升，"
        "再加上前面的业务发展趋势，落地园区之后能带给我们什么经济效益或者产业协同发展的效益。"
        f"（数据：前年营收{fnum(m.revenueY2,0)}万元、去年营收{fnum(m.revenueY1,0)}万元；"
        f"前年税收{fnum(m.taxY2,0)}万元、去年税收{fnum(m.taxY1,0)}万元；行业趋势：{m.industryTrend}。）"
    )

    suggest_header = "建议该项目入园，如果是园中园的项目，就建议该项目以园中园方式入园，一般来说后续的建议主要是这个几点，"
    s1 = "（1）签约  项目招引部门在入驻协议中明确约定企业注册园区，明确经济效益，明确厂房不可转租或分割，做好项目跟踪管理与服务，确保业务按约定导入至新公司，确保项目投产达效。"
    s2 = "（2）产业港项目，产发公司采取“双同步”方式统筹推进项目入驻进度：加快推进产业港项目载体建设，同相关部门做好项目地块的土地征收、规划设计、施工建设、竣工验收等工作，确保产业港具备及时承接能力，为项目顺利落地和企业如期入驻提供有力支撑；加强与企业沟通对接，充分了解其实际需求，加快在园区筛选符合产业规划和企业需求的成熟地块或可用厂房，强化落地时效，防止项目流失或造成“签约不落地”现象。"
    s3 = "（3）如果说企业的经济效益达到了2000万元左右，经济服务局做好项目经济指标跟踪，做好入统指导，帮助企业熟悉统计口径和申报流程，确保企业注册投产达到入统条件后，能够及时、准确地纳入统计范围。"
    s4 = "（4）如果企业能够和园区的主要产业协同发展，互相补位  产发公司与经济服务局协同做好企业服务与培育工作，增强企业发展实力，推动企业与园区协同发展。"
    s5 = "（5）征地项目的话，产发公司与同企业沟通做好土地收储、土地摘牌等相关工作，确保土地收储顺利进行、企业按期进行摘牌建设。产发公司做好项目各项手续服务，协助办理环评、消防、安评等相关手续，确保项目建设和生产合法合规。"

    annex = (
        f"【达标校验】折算亩{fnum(ev['mu'],2)}；固定投资强度{fnum(ev['investIntensity'],1)}万/亩（阈值{ev['thI']}），"
        f"税收强度{fnum(ev['taxIntensity'],1)}万/亩·年（阈值{ev['thT']}）。"
    )
    if not ev["passHard"]:
        if ev["investNeed"] > 0:
            annex += f" 需追加固定投资约{fnum(ev['investNeed'],0)}万元。"
        if ev["taxNeed"] > 0:
            annex += f" 需新增年税收约{fnum(ev['taxNeed'],0)}万元。"

    intent_text = f"是否拟同意入园：{'同意' if m.intentAgree else '不同意'}。"
    decision = f"（系统校验结果：{ev['decision']}）"

    return {
        "intro": intro,
        "judge1": judge1,
        "judge2": judge2,
        "judge3": judge3,
        "suggest_header": suggest_header,
        "s1": s1, "s2": s2, "s3": s3, "s4": s4, "s5": s5,
        "annex": annex,
        "intent": intent_text,
        "decision": decision,
    }

def build_full_report_text(texts: Dict[str,str]) -> str:
    return (
        "一、项目简介\n" + texts["intro"] + "\n\n"
        "二、研判部分\n"
        "1. " + texts["judge1"] + "\n\n"
        "2. " + texts["judge2"] + "\n\n"
        "3. " + texts["judge3"] + "\n\n"
        "三、建议\n" + texts["suggest_header"] + "\n" +
        texts["s1"] + "\n" + texts["s2"] + "\n" + texts["s3"] + "\n" + texts["s4"] + "\n" + texts["s5"] + "\n\n"
        "四、其他\n" + texts["intent"] + "\n" + texts["decision"] + "\n" + texts["annex"]
    )

# ---------------- DeepSeek 接口 ----------------
def _get_client() -> OpenAI | None:
    if "DEEPSEEK_API_KEY" not in st.secrets:
        return None
    return OpenAI(api_key=st.secrets["DEEPSEEK_API_KEY"], base_url="https://api.deepseek.com")

def gen_industry_analysis(m: Model) -> str:
    client = _get_client()
    if client is None:
        return "（未配置 DEEPSEEK_API_KEY，无法生成行业分析）"
    sys = "你是产业园区投资决策的行业分析助手。不要编造具体统计数字或年份；没有把握就说暂无权威公开数据。"
    user = f"""
请基于以下项目信息，给出“行业分析（AI辅助）”一段中文文字（300-600字），按固定4点结构输出：
1）赛道与需求侧
2）供给与竞争
3）区域与协同（石家庄/河北维度）
4）风险点与对策
最后用一句话明确“对本项目落地园区的行业面评价：偏正面/中性/偏谨慎（择一）”。

项目信息：
- 园区：河北石家庄装备制造产业园区
- 产业大类：{INDUSTRIES.get(m.industry,'')}
- 项目类型：{PROJECT_TYPES.get('land','') if m.needType=='buy_land' else PROJECT_TYPES.get('existingNoPolicy','')}
- 拟开展业务：{m.introContent or '—'}
- 企业近两年口径：前年营收{m.revenueY2:.0f}万、去年营收{m.revenueY1:.0f}万；前年税收{m.taxY2:.0f}万、去年税收{m.taxY1:.0f}万
- 行业趋势判断：{m.industryTrend}
"""
    rsp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"system","content":sys},{"role":"user","content":user}],
        temperature=0.2,
        max_tokens=800,
    )
    return rsp.choices[0].message.content.strip()

def gen_polished_report(raw_report: str) -> str:
    client = _get_client()
    if client is None:
        return "（未配置 DEEPSEEK_API_KEY，无法生成润色版）"
    sys = "你是中文公文与园区招商研判报告的润色助手。保持事实与数据不变，不新增数字和结论，不改变阈值口径。"
    user = f"""请在不改变事实数据与阈值口径的前提下，对下面的“园区招商项目研判报告”进行润色：
- 优化语句通顺与层次逻辑
- 保留原有章节标题与顺序（如“一、二、三、四”）
- 不添加外部数据，不改变结论倾向

以下是需要润色的全文：

{raw_report}
"""
    rsp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"system","content":sys},{"role":"user","content":user}],
        temperature=0.2,
        max_tokens=1600,
    )
    return rsp.choices[0].message.content.strip()

# ---------------- 文档导出（严格/润色） ----------------
def build_docx(texts: Dict[str,str]) -> BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)

    doc.add_heading('项目研判报告', level=1)
    doc.add_paragraph('（严格按原文口径生成）')

    doc.add_heading('一、项目简介', level=2)
    doc.add_paragraph(texts["intro"])

    doc.add_heading('二、研判部分', level=2)
    doc.add_paragraph("1. " + texts["judge1"])
    doc.add_paragraph("2. " + texts["judge2"])
    doc.add_paragraph("3. " + texts["judge3"])

    doc.add_heading('三、建议', level=2)
    doc.add_paragraph(texts["suggest_header"])
    for k in ["s1","s2","s3","s4","s5"]:
        doc.add_paragraph(texts[k])

    doc.add_heading('四、其他', level=2)
    doc.add_paragraph(texts["intent"])
    doc.add_paragraph(texts["decision"])
    doc.add_paragraph(texts["annex"])

    # 附录：行业分析（如已生成）
    if "ai_industry" in st.session_state and st.session_state["ai_industry"]:
        doc.add_heading('附：行业分析（AI辅助）', level=2)
        doc.add_paragraph(st.session_state["ai_industry"])

    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio

def build_docx_polished(polished_text: str) -> BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)

    doc.add_heading('项目研判报告（DeepSeek润色版）', level=1)
    for line in polished_text.splitlines():
        if line.strip() == "": 
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio

def build_pdf(texts: Dict[str,str], title: str="项目研判报告") -> BytesIO:
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    W, H = A4
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    c.setTitle(title)

    def draw_text(x, y, text, size=11):
        c.setFont('STSong-Light', size)
        line_height = size * 1.33
        max_chars = 38
        for para in text.split("\\n"):
            while len(para) > max_chars:
                c.drawString(x, y, para[:max_chars]); y -= line_height
                para = para[max_chars:]
            c.drawString(x, y, para); y -= line_height
        return y - 4

    margin = 20*mm; y = H - margin
    c.setFont('STSong-Light', 16); c.drawString(margin, y, title); y -= 10*mm
    c.setFont('STSong-Light', 10); c.drawString(margin, y, "（严格按原文口径生成）" if "润色" not in title else "（AI润色版）"); y -= 8*mm

    c.setFont('STSong-Light', 13); c.drawString(margin, y, "一、项目简介"); y -= 7*mm
    y = draw_text(margin, y, texts["intro"])

    c.setFont('STSong-Light', 13); c.drawString(margin, y, "二、研判部分"); y -= 7*mm
    y = draw_text(margin, y, "1. " + texts["judge1"])
    y = draw_text(margin, y, "2. " + texts["judge2"])
    y = draw_text(margin, y, "3. " + texts["judge3"])

    c.setFont('STSong-Light', 13); c.drawString(margin, y, "三、建议"); y -= 7*mm
    y = draw_text(margin, y, texts["suggest_header"])
    for k in ["s1","s2","s3","s4","s5"]:
        y = draw_text(margin, y, texts[k])

    c.setFont('STSong-Light', 13); c.drawString(margin, y, "四、其他"); y -= 7*mm
    y = draw_text(margin, y, texts["intent"])
    y = draw_text(margin, y, texts["decision"])
    y = draw_text(margin, y, texts["annex"])

    # 附录：行业分析（如已生成）
    if "ai_industry" in st.session_state and st.session_state["ai_industry"]:
        c.setFont('STSong-Light', 13); c.drawString(margin, y, "附：行业分析（AI辅助）"); y -= 7*mm
        y = draw_text(margin, y, st.session_state["ai_industry"])

    c.showPage(); c.save(); bio.seek(0); return bio

def build_pdf_polished(polished_text: str) -> BytesIO:
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    W, H = A4
    bio = BytesIO(); c = canvas.Canvas(bio, pagesize=A4); c.setTitle("项目研判报告（DeepSeek润色版）")
    def draw_text(x, y, text, size=11):
        c.setFont('STSong-Light', size); line_height = size * 1.33; max_chars = 38
        for para in text.split("\\n"):
            while len(para) > max_chars:
                c.drawString(x, y, para[:max_chars]); y -= line_height; para = para[max_chars:]
            c.drawString(x, y, para); y -= line_height
        return y - 4
    margin=20*mm; y=A4[1]-margin
    c.setFont('STSong-Light', 16); c.drawString(margin, y, "项目研判报告（DeepSeek润色版）"); y -= 10*mm
    c.setFont('STSong-Light', 10); c.drawString(margin, y, "（AI润色版）"); y -= 8*mm
    y = draw_text(margin, y, polished_text)
    c.showPage(); c.save(); bio.seek(0); return bio

# ---------------- UI ----------------
st.markdown("## 石家庄装备制造产业园 — 投资决策委员会 项目研判（严格口径 + DeepSeek 行业分析/润色）")

with st.sidebar:
    st.subheader("阈值设置")
    CONFIG["thresholds"]["investPerMu"] = float(st.number_input("固定投资阈值（万/亩）", 0.0, 10000.0, CONFIG["thresholds"]["investPerMu"]))
    CONFIG["thresholds"]["taxPerMu"] = float(st.number_input("税收强度阈值（万/亩·年）", 0.0, 1000.0, CONFIG["thresholds"]["taxPerMu"]))

c1, c2 = st.columns(2)
with c1:
    st.header("一、项目简介（输入）")
    projectName = st.text_input("项目名称", value="高端装备制造项目")
    investTotal = st.number_input("计划投资（万元）", min_value=0.0, value=10000.0, step=100.0)
    locate = st.text_input("拟选址", value="栾城区科创谷/产业港")
    landMu = st.number_input("占地（亩）", min_value=0.0, value=0.0, step=0.01)
    buildingArea = st.number_input("实际建筑面积（平米）", min_value=0.0, value=0.0, step=1.0)
    floorRatio = st.number_input("容积率（折算用）", min_value=0.0, value=0.0, step=0.1)
    introContent = st.text_area("建设内容", value="新建高精度机加工产线与装配线")
    expectedOutput = st.number_input("预计年产值（万元）", min_value=0.0, value=20000.0, step=100.0)
    expectedAnnualTax = st.number_input("预计年税收（万元）", min_value=0.0, value=1500.0, step=10.0)
    expectedJobs = st.number_input("预计带动就业人数（人）", min_value=0, value=200, step=1)

with c2:
    st.header("二、研判要点（输入）")
    industry = st.selectbox("产业大类", options=list(INDUSTRIES.keys()), format_func=lambda k: INDUSTRIES[k])
    companyName = st.text_input("项目主体（公司名称）", value="某装备制造有限公司")
    establishedYear = st.text_input("注册年份", value="2016")
    registeredAt = st.text_input("注册地", value="石家庄市栾城区")
    isLuanReg = st.checkbox("注册在石家庄市栾城区（否则视为尚未注册于园区）", value=True)
    importBusiness = st.text_input("拟将导入园区的业务", value="核心零部件制造与总装")
    newBusiness = st.text_input("扩展新业务（可选）", value="")

    needType = st.selectbox("主要需求", options=list(NEED_TYPES.keys()), format_func=lambda k: NEED_TYPES[k])
    carrier = st.selectbox("拟承接载体", options=list(CARRIERS.keys()), format_func=lambda k: CARRIERS[k])
    techTitles = st.text_input("技术/称号（河北省/国家级等）", value="省级专精特新、高新技术企业")
    chainMaturity = st.selectbox("产业链条成熟度", options=["完善","成熟","一般"], index=1)
    innovation = st.selectbox("技术创新能力", options=["强","较强","一般"], index=1)
    customerStable = st.selectbox("客户资源稳定性", options=["稳定","一般","不稳定"], index=0)
    marketBase = st.selectbox("市场基础", options=["扎实","一般","较弱"], index=0)
    chainSegmentFill = st.text_input("入园后可填补的产业链环节（可选）", value="关键零部件加工")

    st.header("三、经营与趋势（输入）")
    revenueY2 = st.number_input("前年营收（万元）", min_value=0.0, value=18000.0, step=100.0)
    revenueY1 = st.number_input("去年营收（万元）", min_value=0.0, value=22000.0, step=100.0)
    taxY2 = st.number_input("前年税收（万元）", min_value=0.0, value=1100.0, step=10.0)
    taxY1 = st.number_input("去年税收（万元）", min_value=0.0, value=1300.0, step=10.0)
    industryTrend = st.selectbox("行业趋势", options=["向好","平稳","承压"], index=0)

    st.header("四、风险与意向")
    riskDishonest = st.checkbox("失信被执行/严重信用风险", value=False)
    riskEnv = st.checkbox("重大环保/安监处罚未结", value=False)
    riskIllegalLand = st.checkbox("违法违规用地", value=False)
    riskLicenseMissing = st.checkbox("核心资质缺失且短期不可补齐", value=False)
    intentAgree = st.checkbox("是否拟同意入园", value=True)

# 组装与评估
m = Model(
    projectName=projectName, investTotal=investTotal, locate=locate, landMu=landMu,
    buildingArea=buildingArea, floorRatio=floorRatio, introContent=introContent,
    expectedOutput=expectedOutput, expectedAnnualTax=expectedAnnualTax, expectedJobs=int(expectedJobs),
    industry=industry, companyName=companyName, establishedYear=establishedYear, registeredAt=registeredAt,
    isLuanReg=isLuanReg, importBusiness=importBusiness, newBusiness=newBusiness,
    needType=needType, carrier=carrier, techTitles=techTitles, chainMaturity=chainMaturity,
    innovation=innovation, customerStable=customerStable, marketBase=marketBase, chainSegmentFill=chainSegmentFill,
    revenueY2=revenueY2, revenueY1=revenueY1, taxY2=taxY2, taxY1=taxY1, industryTrend=industryTrend,
    riskDishonest=riskDishonest, riskEnv=riskEnv, riskIllegalLand=riskIllegalLand, riskLicenseMissing=riskLicenseMissing,
    intentAgree=intentAgree,
)
ev = evaluate(m)
texts = build_text(m, ev)

# 预览（严格口径）
st.markdown("---")
st.markdown("### 预览 · 研判报告（严格按原文）")
st.markdown(f"""
**一、项目简介**  
{texts['intro']}

**二、研判部分**  
1. {texts['judge1']}

2. {texts['judge2']}

3. {texts['judge3']}

**三、建议**  
{texts['suggest_header']}
{texts['s1']}
{texts['s2']}
{texts['s3']}
{texts['s4']}
{texts['s5']}

**四、其他**  
{texts['intent']}  
{texts['decision']}  
{texts['annex']}
""")

# DeepSeek：行业分析（直接体现在页面）
def _get_client() -> OpenAI | None:
    if "DEEPSEEK_API_KEY" not in st.secrets:
        return None
    return OpenAI(api_key=st.secrets["DEEPSEEK_API_KEY"], base_url="https://api.deepseek.com")

def gen_industry_analysis(m: Model) -> str:
    client = _get_client()
    if client is None:
        return "（未配置 DEEPSEEK_API_KEY，无法生成行业分析）"
    sys = "你是产业园区投资决策的行业分析助手。不要编造具体统计数字或年份；没有把握就说暂无权威公开数据。"
    user = f"""
请基于以下项目信息，给出“行业分析（AI辅助）”一段中文文字（300-600字），按固定4点结构输出：
1）赛道与需求侧
2）供给与竞争
3）区域与协同（石家庄/河北维度）
4）风险点与对策
最后用一句话明确“对本项目落地园区的行业面评价：偏正面/中性/偏谨慎（择一）”。

项目信息：
- 园区：河北石家庄装备制造产业园区
- 产业大类：{INDUSTRIES.get(m.industry,'')}
- 项目类型：{PROJECT_TYPES.get('land','') if m.needType=='buy_land' else PROJECT_TYPES.get('existingNoPolicy','')}
- 拟开展业务：{m.introContent or '—'}
- 企业近两年口径：前年营收{m.revenueY2:.0f}万、去年营收{m.revenueY1:.0f}万；前年税收{m.taxY2:.0f}万、去年税收{m.taxY1:.0f}万
- 行业趋势判断：{m.industryTrend}
"""
    rsp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"system","content":sys},{"role":"user","content":user}],
        temperature=0.2,
        max_tokens=800,
    )
    return rsp.choices[0].message.content.strip()

def gen_polished_report(raw_report: str) -> str:
    client = _get_client()
    if client is None:
        return "（未配置 DEEPSEEK_API_KEY，无法生成润色版）"
    sys = "你是中文公文与园区招商研判报告的润色助手。保持事实与数据不变，不新增数字和结论，不改变阈值口径。"
    user = f"""请在不改变事实数据与阈值口径的前提下，对下面的“园区招商项目研判报告”进行润色：
- 优化语句通顺与层次逻辑
- 保留原有章节标题与顺序（如“一、二、三、四”）
- 不添加外部数据，不改变结论倾向

以下是需要润色的全文：

{raw_report}
"""
    rsp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"system","content":sys},{"role":"user","content":user}],
        temperature=0.2,
        max_tokens=1600,
    )
    return rsp.choices[0].message.content.strip()

with st.expander("附：行业分析（AI辅助 · DeepSeek）", expanded=False):
    if st.button("生成行业分析", use_container_width=True):
        with st.spinner("DeepSeek 正在生成行业分析…"):
            st.session_state["ai_industry"] = gen_industry_analysis(m)
    if "ai_industry" in st.session_state and st.session_state["ai_industry"]:
        st.write(st.session_state["ai_industry"])

with st.expander("DEEPSEEK 润色版（对上面整份报告进行润色）", expanded=False):
    if st.button("生成润色版", use_container_width=True):
        raw_text = (
            "一、项目简介\n" + texts["intro"] + "\n\n"
            "二、研判部分\n"
            "1. " + texts["judge1"] + "\n\n"
            "2. " + texts["judge2"] + "\n\n"
            "3. " + texts["judge3"] + "\n\n"
            "三、建议\n" + texts["suggest_header"] + "\n" +
            texts["s1"] + "\n" + texts["s2"] + "\n" + texts["s3"] + "\n" + texts["s4"] + "\n" + texts["s5"] + "\n\n"
            "四、其他\n" + texts["intent"] + "\n" + texts["decision"] + "\n" + texts["annex"]
        )
        if "ai_industry" in st.session_state and st.session_state["ai_industry"]:
            raw_text += "\n\n附：行业分析（AI辅助）\n" + st.session_state["ai_industry"]
        with st.spinner("DeepSeek 正在润色…"):
            st.session_state["ai_polished"] = gen_polished_report(raw_text)
    if "ai_polished" in st.session_state and st.session_state["ai_polished"]:
        st.markdown("**润色版预览：**")
        st.write(st.session_state["ai_polished"])

# 导出按钮区
def build_docx(texts: Dict[str,str]) -> BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)

    doc.add_heading('项目研判报告', level=1)
    doc.add_paragraph('（严格按原文口径生成）')

    doc.add_heading('一、项目简介', level=2)
    doc.add_paragraph(texts["intro"])

    doc.add_heading('二、研判部分', level=2)
    doc.add_paragraph("1. " + texts["judge1"])
    doc.add_paragraph("2. " + texts["judge2"])
    doc.add_paragraph("3. " + texts["judge3"])

    doc.add_heading('三、建议', level=2)
    doc.add_paragraph(texts["suggest_header"])
    for k in ["s1","s2","s3","s4","s5"]:
        doc.add_paragraph(texts[k])

    doc.add_heading('四、其他', level=2)
    doc.add_paragraph(texts["intent"])
    doc.add_paragraph(texts["decision"])
    doc.add_paragraph(texts["annex"])

    if "ai_industry" in st.session_state and st.session_state["ai_industry"]:
        doc.add_heading('附：行业分析（AI辅助）', level=2)
        doc.add_paragraph(st.session_state["ai_industry"])

    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio

def build_pdf(texts: Dict[str,str], title: str="项目研判报告") -> BytesIO:
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    W, H = A4; bio = BytesIO(); c = canvas.Canvas(bio, pagesize=A4); c.setTitle(title)
    def draw_text(x, y, text, size=11):
        c.setFont('STSong-Light', size); line_height = size * 1.33; max_chars = 38
        for para in text.split("\\n"):
            while len(para) > max_chars:
                c.drawString(x, y, para[:max_chars]); y -= line_height; para = para[max_chars:]
            c.drawString(x, y, para); y -= line_height
        return y - 4
    margin=20*mm; y=A4[1]-margin
    c.setFont('STSong-Light', 16); c.drawString(margin, y, title); y -= 10*mm
    c.setFont('STSong-Light', 10); c.drawString(margin, y, "（严格按原文口径生成）" if "润色" not in title else "（AI润色版）"); y -= 8*mm
    c.setFont('STSong-Light', 13); c.drawString(margin, y, "一、项目简介"); y -= 7*mm
    y = draw_text(margin, y, texts["intro"])
    c.setFont('STSong-Light', 13); c.drawString(margin, y, "二、研判部分"); y -= 7*mm
    y = draw_text(margin, y, "1. " + texts["judge1"])
    y = draw_text(margin, y, "2. " + texts["judge2"])
    y = draw_text(margin, y, "3. " + texts["judge3"])
    c.setFont('STSong-Light', 13); c.drawString(margin, y, "三、建议"); y -= 7*mm
    y = draw_text(margin, y, texts["suggest_header"])
    for k in ["s1","s2","s3","s4","s5"]:
        y = draw_text(margin, y, texts[k])
    c.setFont('STSong-Light', 13); c.drawString(margin, y, "四、其他"); y -= 7*mm
    y = draw_text(margin, y, texts["intent"]); y = draw_text(margin, y, texts["decision"]); y = draw_text(margin, y, texts["annex"])
    if "ai_industry" in st.session_state and st.session_state["ai_industry"]:
        c.setFont('STSong-Light', 13); c.drawString(margin, y, "附：行业分析（AI辅助）"); y -= 7*mm
        y = draw_text(margin, y, st.session_state["ai_industry"])
    c.showPage(); c.save(); bio.seek(0); return bio

docx_bytes = build_docx(texts)
pdf_bytes = build_pdf(texts, title="项目研判报告")
c1, c2 = st.columns(2)
c1.download_button("导出 Word（严格口径）", data=docx_bytes, file_name=f"{projectName or '项目'}_研判报告_严格口径.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
c2.download_button("导出 PDF（严格口径）", data=pdf_bytes, file_name=f"{projectName or '项目'}_研判报告_严格口径.pdf", mime="application/pdf")

if "ai_polished" in st.session_state and st.session_state["ai_polished"]:
    from io import BytesIO as _BytesIO
    def build_docx_polished(polished_text: str) -> _BytesIO:
        doc = Document(); style = doc.styles['Normal']; style.font.name='宋体'; style._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); style.font.size=Pt(11)
        doc.add_heading('项目研判报告（DeepSeek润色版）', level=1)
        for line in polished_text.splitlines():
            doc.add_paragraph(line if line.strip() else "")
        bio = _BytesIO(); doc.save(bio); bio.seek(0); return bio
    def build_pdf_polished(polished_text: str) -> _BytesIO:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light')); W,H=A4; bio=_BytesIO(); c=canvas.Canvas(bio,pagesize=A4); c.setTitle("项目研判报告（DeepSeek润色版）")
        def draw_text(x, y, text, size=11):
            c.setFont('STSong-Light', size); line_height=size*1.33; max_chars=38
            for para in text.split("\\n"):
                while len(para) > max_chars:
                    c.drawString(x, y, para[:max_chars]); y -= line_height; para = para[max_chars:]
                c.drawString(x, y, para); y -= line_height
            return y-4
        margin=20*mm; y=A4[1]-margin
        c.setFont('STSong-Light', 16); c.drawString(margin, y, "项目研判报告（DeepSeek润色版）"); y -= 10*mm
        c.setFont('STSong-Light', 10); c.drawString(margin, y, "（AI润色版）"); y -= 8*mm
        y = draw_text(margin, y, st.session_state["ai_polished"]); c.showPage(); c.save(); bio.seek(0); return bio
    pol_docx = build_docx_polished(st.session_state["ai_polished"])
    pol_pdf = build_pdf_polished(st.session_state["ai_polished"])
    p1, p2 = st.columns(2)
    p1.download_button("导出 Word（DeepSeek润色版）", data=pol_docx, file_name=f"{projectName or '项目'}_研判报告_DeepSeek润色版.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    p2.download_button("导出 PDF（DeepSeek润色版）", data=pol_pdf, file_name=f"{projectName or '项目'}_研判报告_DeepSeek润色版.pdf", mime="application/pdf")
